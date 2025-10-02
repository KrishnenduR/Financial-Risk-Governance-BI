#!/usr/bin/env python3
"""
Extract content from the Project Structure and Word Distribution DOCX file
to understand the required documentation structure.
"""

import os
from pathlib import Path
from docx import Document

def extract_docx_content(file_path):
    """Extract all text content from a DOCX file."""
    try:
        doc = Document(file_path)
        content = []
        
        print(f"📄 Reading DOCX file: {file_path}")
        print("="*80)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # Check if it's a heading
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    print(f"{'#' * int(level) if level.isdigit() else '#'} {text}")
                    content.append(f"{'#' * int(level) if level.isdigit() else '#'} {text}")
                else:
                    print(text)
                    content.append(text)
                print()  # Empty line for readability
        
        # Also extract tables if any
        print("\n" + "="*80)
        print("📊 TABLES FOUND:")
        print("="*80)
        
        for i, table in enumerate(doc.tables):
            print(f"\n--- TABLE {i+1} ---")
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                print(" | ".join(row_data))
        
        return content
    
    except Exception as e:
        print(f"❌ Error reading DOCX file: {e}")
        return None

def main():
    """Main function to extract DOCX content."""
    docx_path = "Project Structure and Word Distribution.docx"
    
    if not os.path.exists(docx_path):
        print(f"❌ File not found: {docx_path}")
        return
    
    content = extract_docx_content(docx_path)
    
    if content:
        # Save extracted content to a text file for easier access
        with open("extracted_project_structure.txt", "w", encoding="utf-8") as f:
            f.write("\n".join(content))
        print(f"\n✅ Content extracted and saved to 'extracted_project_structure.txt'")
    else:
        print("❌ Failed to extract content")

if __name__ == "__main__":
    main()